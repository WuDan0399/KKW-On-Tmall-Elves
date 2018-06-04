import pymysql
import docx

class Connection(object):
    def connect_database(self):
        self.connection = pymysql.connect(host='localhost',
                                          user='root',
                                          db='kaokaoni',
                                          passwd='000000',
                                          charset="utf8"
                                          )
        self.cursor = self.connection.cursor()

    def disconnect_database(self):
        self.cursor.close()
        self.connection.close()

    def exec_query(self, sql):
        self.cursor.execute(sql)

    def exec_update(self, sql):
        self.cursor.execute(sql)
        self.connection.commit()

    def fetch_cursor(self):
        return self.cursor.fetchall()

file = docx.Document('电信类.docx')
i = 0
qa_pair = dict()
qr_pair = dict()
q_list = list()
while i < len(file.paragraphs):
    q_list.append(file.paragraphs[i].text)
    qa_pair[file.paragraphs[i].text] = file.paragraphs[i+1].text
    qr_pair[file.paragraphs[i].text] = file.paragraphs[i+2].text
    i = i+4
database = Connection()
database.connect_database()
for i in range(len(q_list)):
        database.exec_query("""INSERT INTO questions VALUES(%d,"%s","%s","%s","%s")"""%(i+1, q_list[i], qa_pair[q_list[i]], qr_pair[q_list[i]],'电信类'))
database.disconnect_database()
